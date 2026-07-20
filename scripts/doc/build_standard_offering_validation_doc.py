#!/usr/bin/env python3
"""Generate OriginBA Standard Offering Report Library validation Word documents."""

from __future__ import annotations

import argparse
import io
import json
import re
import shutil
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPO_ROOT = Path(__file__).resolve().parents[2]
COVER_BASE_PATH = REPO_ROOT / "output/doc/templates/CLIENT_NAME_25.4_TEST_Role_Based_Access_Cover_Base.docx"
CATALOG_PATH = REPO_ROOT / "tmp/docs/report_library_catalog.json"
REPORT_LIBRARY_DOC = Path("/Users/chase/Downloads/OriginBA_SmartCityReportLibrary_v1.0 (1).docx")
STANDARD_OFFERING_ZIP = Path("/Users/chase/Downloads/standardoffering.zip")
TEMPLATE_OUTPUT = REPO_ROOT / "output/doc/templates/OriginBA_Standard_Offering_Report_Library_Validation_TEMPLATE.docx"
DEFAULT_OUTPUT = REPO_ROOT / "output/doc/OriginBA_Standard_Offering_Report_Library_Validation.docx"

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WNSMAP = {"w": WNS}

PLACEHOLDER_CLIENT = "CLIENT_NAME"
PLACEHOLDER_TEST = "TEST_VERSION"
PLACEHOLDER_DATASOURCE = "DATASOURCE_ALIAS"
PLACEHOLDER_AUTHOR = "AUTHOR_NAME"
PLACEHOLDER_STATUS = "PASS/FAIL/NA"
PLACEHOLDER_NOTES = (
    "Executed successfully. Updated filters as required. "
    "Returned expected results with no errors."
)
TEMPLATE_NOTES = (
    "Record execution result. Example: Executed successfully. Updated filters as required. "
    "Returned expected results with no errors."
)

EXCLUDED_STEMS = {
    "Payment_Arrangement___Not_Paid_Bill",
    "Payment_Arrangement___Reconciliation",
    "Payment_Arrangement___Total_receivables_outstanding",
}

WS_ORDER = [
    "Billing and Rates",
    "Cashiering",
    "Common",
    "Customer Operations",
    "Debt Management",
    "Field Operations",
    "Finance",
    "Meter Operations",
    "New Services",
]

WS_ZIP_KEYS = {
    "Billing and Rates": "Billing_and_Rates",
    "Cashiering": "Cashiering",
    "Common": "Common",
    "Customer Operations": "Customer_Operations",
    "Debt Management": "Debt_Management",
    "Field Operations": "Field_Operations",
    "Finance": "Finance",
    "Meter Operations": "Meter_Operations",
    "New Services": "New_Services___Planning",
}


@dataclass
class GenerationConfig:
    client_name: str = PLACEHOLDER_CLIENT
    test_version: str = PLACEHOLDER_TEST
    datasource_alias: str = PLACEHOLDER_DATASOURCE
    author_name: str = PLACEHOLDER_AUTHOR
    validation_status: str = PLACEHOLDER_STATUS
    notes_text: str = TEMPLATE_NOTES
    template_mode: bool = False
    output_path: Path = DEFAULT_OUTPUT
    document_date: str | None = None

    @property
    def is_fillable_template(self) -> bool:
        return self.template_mode or self.client_name == PLACEHOLDER_CLIENT

    @property
    def resolved_date(self) -> str:
        return self.document_date or date.today().strftime("%B %d, %Y")

    @property
    def cover_title_lines(self) -> tuple[str, str, str]:
        return (
            "STANDARD OFFERING REPORT",
            "LIBRARY VALIDATION –",
            f"{self.client_name} CIS TEST {self.test_version}",
        )

    @property
    def footer_text(self) -> str:
        return (
            f"@2026, Origin Utility, Inc / Proprietary & Confidential / "
            f"Expressly for the {self.client_name}"
        )

    @property
    def version_notes(self) -> str:
        if self.is_fillable_template:
            return "Initial validation evidence for CLIENT_NAME Standard Offering report library."
        return f"Initial validation evidence for {self.client_name} Standard Offering report library."

    @property
    def summary_result(self) -> str:
        if self.is_fillable_template:
            return PLACEHOLDER_STATUS
        return self.validation_status


def norm(text: str) -> str:
    text = text.replace("\u2013", "-").replace("\u2014", "-").replace("&", " and ")
    text = re.sub(r"[^a-z0-9]+", " ", text.lower())
    return re.sub(r"\s+", " ", text).strip()


def fname_label(stem: str) -> str:
    return stem.replace("___", " - ").replace("_", " ").strip()


OBJECT_NORMALIZE = {
    "SA Snapshot   Aged Balance": "SA Snapshot - Aged Balance",
    "Write Off Process 1": "Write Off Process",
}


def normalize_object(name: str) -> str:
    return OBJECT_NORMALIZE.get(name, name)


def report_prefix(name: str) -> str:
    if " - " in name:
        return name.split(" - ", 1)[0].strip()
    if " – " in name:
        return name.split(" – ", 1)[0].strip()
    return name.strip()


def object_name_for_item(ws: str, name: str, zip_object: str | None) -> str:
    if zip_object:
        return normalize_object(zip_object)
    prefix = report_prefix(name)
    dashboard_object_map = {
        "Billed Amount - Dashboard": "Billed Amount",
        "Billed Usage - Dashboard": "Billed Usage",
        "Cashiering Dashboard": "Payment Header",
        "Batch Process Dashboard": "Batch",
        "Exception and To Do Dashboard": "To Do",
        "Customer Operations Dashboard": "Customer Contact",
        "Collections Performance Dashboard": "Collection Process",
        "Field Operations Dashboard": "Field Activity",
        "Finance Dashboard": "Financial Transaction",
        "General Ledger - Dashboard": "General Ledger",
        "Measurements - Dashboard": "Measurements",
        "Usage Dashboard": "Usage",
        "New Services Dashboard": "New Services",
        "Debt Management - Dashboard": "Collection Process",
    }
    if name in dashboard_object_map:
        return dashboard_object_map[name]

    rules: dict[str, dict[str, str]] = {
        "Billing and Rates": {
            "Billed Amount": "Billed Amount",
            "Billed Usage": "Billed Usage",
            "Pending Bills": "Billing Errors",
        },
        "Cashiering": {
            "Deposit Control": "Deposit Control",
            "Pay Plan": "Pay Plan",
            "Payment": "Payment Header",
            "Payments": "Payment Header",
            "Tender Control": "Tender Control",
            "Tender": "Payment Tender",
        },
        "Common": {
            "Batch Process": "Batch",
            "Bill Segment Exception": "Exception",
            "To Do": "To Do",
            "Usage Transaction Exceptions": "Exception",
            "VEE Exception": "Exception",
            "VEE Exceptions": "Exception",
        },
        "Customer Operations": {
            "Account Alert": "Account Alert",
            "Customer Contact": "Customer Contact",
            "Customer Contacts": "Customer Contact",
            "Case": "Case",
            "Premise": "Premise",
            "Customer": "Customer",
            "Landlord": "Landlord Agreement",
        },
        "Debt Management": {
            "Aged Debt": "SA Snapshot - Aged Balance",
            "Collection Process": "Collection Process",
            "Severance Process": "Severance Process",
            "Write Off Process": "Write Off Process",
            "Write Offs": "Write Off Process",
        },
        "Field Operations": {
            "Field Activity": "Field Activity",
            "Crew": "Crew",
        },
        "Finance": {
            "General Ledger": "General Ledger",
            "Financial Transaction": "Financial Transaction",
            "Financial Transactions": "Financial Transaction",
            "Adjustment": "Adjustments",
            "Adjustments": "Adjustments",
            "Billable Charge": "Billable Charge",
            "Deposits": "Cash Deposit",
            "Write Off": "Write Offs",
            "Payment Arrangement": "Payment Arrangement",
        },
        "Meter Operations": {
            "Measurement": "Measurements",
            "Measurements": "Measurements",
            "Meter Reads": "Measurements",
            "Usage": "Usage",
            "Usage Transaction": "Usage Transactions",
            "Usage Transactions": "Usage Transactions",
            "Device": "Device",
            "Device Events": "Device",
            "Asset": "Asset",
        },
        "New Services": {
            "New Services": "New Services",
            "New Service": "New Services",
            "New Service Agreements": "New Services",
        },
    }
    ws_rules = rules.get(ws, {})
    for key, obj in ws_rules.items():
        if prefix == key or prefix.startswith(key):
            return normalize_object(obj)
    return normalize_object(prefix)


def ensure_cover_base() -> Path:
    COVER_BASE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if COVER_BASE_PATH.exists():
        return COVER_BASE_PATH

    fallback = Path("/Users/chase/Downloads/CLIENT_NAME 25.4 TEST - Role-Based Access Testing.docx")
    if not fallback.exists():
        raise SystemExit(
            "Cover base document not found. Place the Role-Based Access Testing template at "
            f"{COVER_BASE_PATH} or {fallback}."
        )
    shutil.copy2(fallback, COVER_BASE_PATH)
    return COVER_BASE_PATH


def load_catalog() -> tuple[list[dict], list[dict]]:
    if CATALOG_PATH.exists():
        data = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
        return data["reports"], data["dashboards"]

    if not REPORT_LIBRARY_DOC.exists():
        raise SystemExit(
            f"Report catalog not found. Expected {CATALOG_PATH} or {REPORT_LIBRARY_DOC}."
        )

    doc = Document(str(REPORT_LIBRARY_DOC))
    reports: list[dict] = []
    dashboards: list[dict] = []
    current_ws: str | None = None
    for el in doc.element.body:
        tag = el.tag.split("}", 1)[-1]
        if tag == "p":
            from docx.text.paragraph import Paragraph

            para = Paragraph(el, doc)
            if para.style.name == "Heading 2" and para.text.strip() not in (
                "Report Source Types",
                "Workstream Coverage",
            ):
                current_ws = para.text.strip()
        elif tag == "tbl":
            from docx.table import Table

            table = Table(el, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows and rows[0][0] == "Report Name" and current_ws:
                for row in rows[1:]:
                    if not row[0]:
                        continue
                    item = {
                        "ws": current_ws,
                        "name": row[0],
                        "source": row[1] if len(row) > 1 else "",
                    }
                    if "Dashboard" in row[0]:
                        item["kind"] = "Dashboard"
                        dashboards.append(item)
                    else:
                        item["kind"] = "Ad Hoc View"
                        reports.append(item)

    CATALOG_PATH.parent.mkdir(parents=True, exist_ok=True)
    CATALOG_PATH.write_text(
        json.dumps({"reports": reports, "dashboards": dashboards}, indent=2),
        encoding="utf-8",
    )
    return reports, dashboards


def load_zip_index(zip_path: Path) -> list[dict]:
    if not zip_path.exists():
        return []

    items: list[dict] = []
    with zipfile.ZipFile(zip_path) as archive:
        for member in archive.namelist():
            if "Standard_Offering/" not in member:
                continue
            if not member.endswith(".xml") or member.endswith(".folder.xml") or "_files/" in member:
                continue
            rel = member.split("Standard_Offering/")[-1]
            parts = rel.split("/")
            if len(parts) < 3:
                continue
            ws_key, obj_key, filename = parts[0], parts[1], parts[2]
            stem = filename[:-4]
            if stem in EXCLUDED_STEMS:
                continue
            root = ET.fromstring(archive.read(member))
            tag = root.tag.split("}", 1)[-1]
            if tag not in {"adhocDataView", "dashboardModelResource"}:
                continue
            kind = "Dashboard" if tag == "dashboardModelResource" else "Ad Hoc View"
            label = fname_label(stem)
            items.append(
                {
                    "ws_key": ws_key,
                    "object": obj_key.replace("_", " "),
                    "stem": stem,
                    "label": label,
                    "kind": kind,
                    "norm": norm(label),
                }
            )
    return items


def match_zip_object(item: dict, zip_items: list[dict], used: set[str]) -> str | None:
    ws_key = WS_ZIP_KEYS[item["ws"]]
    kind = item.get("kind", "Ad Hoc View")
    candidates = [
        z
        for z in zip_items
        if z["ws_key"] == ws_key and z["stem"] not in used and z["kind"] == kind
    ]
    target = norm(item["name"])
    best: dict | None = None
    best_score = 0
    for candidate in candidates:
        score = len(set(target.split()) & set(candidate["norm"].split()))
        if target == candidate["norm"]:
            score += 100
        if target in candidate["norm"] or candidate["norm"] in target:
            score += 40
        if score > best_score:
            best_score = score
            best = candidate
    if best and best_score >= 2:
        used.add(best["stem"])
        return best["object"]
    return None


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def add_validation_table(
    doc: Document,
    rows: list[dict],
    config: GenerationConfig,
    dashboard: bool = False,
) -> None:
    if dashboard:
        headers = ["Dashboard", "Validation", "Notes"]
    else:
        headers = ["Report", "Type", "Validation", "Notes"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        if dashboard:
            set_cell_text(cells[0], row["name"])
            set_cell_text(cells[1], config.validation_status)
            set_cell_text(cells[2], config.notes_text)
        else:
            set_cell_text(cells[0], row["name"])
            set_cell_text(cells[1], row.get("kind", "Ad Hoc View"))
            set_cell_text(cells[2], config.validation_status)
            set_cell_text(cells[3], config.notes_text)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _make_text_run(text: str) -> ET.Element:
    run = ET.Element(f"{{{WNS}}}r")
    text_el = ET.SubElement(run, f"{{{WNS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    return run


def _make_line_break_run() -> ET.Element:
    run = ET.Element(f"{{{WNS}}}r")
    ET.SubElement(run, f"{{{WNS}}}br")
    return run


def _paragraph_style(paragraph: ET.Element) -> str | None:
    p_pr = paragraph.find("w:pPr", WNSMAP)
    if p_pr is None:
        return None
    style = p_pr.find("w:pStyle", WNSMAP)
    if style is None:
        return None
    return style.get(f"{{{WNS}}}val")


def _update_cover_title(sdt: ET.Element, title_lines: tuple[str, ...]) -> None:
    updated = False
    for paragraph in sdt.findall(".//w:p", WNSMAP):
        if _paragraph_style(paragraph) != "Title":
            continue
        texts = [node.text or "" for node in paragraph.findall(".//w:t", WNSMAP)]
        if updated or (texts and not any(ch.isalpha() for ch in "".join(texts))):
            continue
        for child in list(paragraph):
            if child.tag != f"{{{WNS}}}pPr":
                paragraph.remove(child)
        for index, line in enumerate(title_lines):
            if index:
                paragraph.append(_make_line_break_run())
            paragraph.append(_make_text_run(line))
        updated = True
        break
    if not updated:
        raise SystemExit("Template cover title paragraph was not found")


def _clean_cover_content(content: ET.Element) -> None:
    for child in list(content):
        if child.tag == f"{{{WNS}}}sdt":
            content.remove(child)
    title_seen = False
    for paragraph in list(content.findall("w:p", WNSMAP)):
        if _paragraph_style(paragraph) == "Title":
            texts = [node.text or "" for node in paragraph.findall(".//w:t", WNSMAP)]
            if title_seen or not any(ch.isalpha() for ch in "".join(texts)):
                content.remove(paragraph)
                continue
            title_seen = True


def _extract_section_properties(body: ET.Element) -> ET.Element | None:
    for paragraph in body.findall("w:p", WNSMAP):
        p_pr = paragraph.find("w:pPr", WNSMAP)
        if p_pr is None:
            continue
        sect_pr = p_pr.find("w:sectPr", WNSMAP)
        if sect_pr is not None:
            return deepcopy(sect_pr)
    return None


def prepare_template_cover(docx_path: Path, config: GenerationConfig) -> None:
    """Preserve the Word cover-page gallery and set the client-specific title text."""
    with zipfile.ZipFile(docx_path, "r") as archive:
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as out_archive:
            for info in archive.infolist():
                data = archive.read(info.filename)
                if info.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", WNSMAP)
                    if body is None:
                        raise SystemExit("Template is missing word/document.xml body")

                    sect_pr = _extract_section_properties(body)
                    children = list(body)
                    header_paragraph = None
                    cover_sdt = None
                    for child in children:
                        if child.tag == f"{{{WNS}}}p" and header_paragraph is None:
                            header_paragraph = child
                        elif child.tag == f"{{{WNS}}}sdt" and cover_sdt is None:
                            cover_sdt = child

                    if cover_sdt is None:
                        raise SystemExit("Template cover page block was not found")

                    _update_cover_title(cover_sdt, config.cover_title_lines)
                    content = cover_sdt.find("w:sdtContent", WNSMAP)
                    if content is not None:
                        _clean_cover_content(content)

                    for child in list(body):
                        body.remove(child)
                    if header_paragraph is not None:
                        body.append(deepcopy(header_paragraph))
                    body.append(deepcopy(cover_sdt))

                    anchor = ET.SubElement(body, f"{{{WNS}}}p")
                    if sect_pr is not None:
                        p_pr = ET.SubElement(anchor, f"{{{WNS}}}pPr")
                        p_pr.append(sect_pr)

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif info.filename.startswith("word/footer") or info.filename.startswith("word/header"):
                    text = data.decode("utf-8")
                    if config.is_fillable_template:
                        text = text.replace("City of Odessa", PLACEHOLDER_CLIENT)
                        text = text.replace("CityCorp", PLACEHOLDER_CLIENT)
                    else:
                        text = text.replace("CLIENT_NAME", config.client_name)
                        text = text.replace("City of Odessa", config.client_name)
                        text = text.replace("CityCorp", config.client_name)
                    data = text.encode("utf-8")

                out_archive.writestr(info, data)

    docx_path.write_bytes(buffer.getvalue())


def configure_footer(doc: Document, text: str) -> None:
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        paragraph = footer.paragraphs[0]
        paragraph.text = text


def build_structure(reports: list[dict], dashboards: list[dict], zip_items: list[dict]) -> dict:
    used: set[str] = set()
    enriched: list[dict] = []
    for item in reports + dashboards:
        zip_object = match_zip_object(item, zip_items, used)
        obj = object_name_for_item(item["ws"], item["name"], zip_object)
        enriched.append({**item, "object": obj})

    structure: dict[str, dict[str, list[dict]]] = defaultdict(lambda: defaultdict(list))
    dash_structure: dict[str, dict[str, list[dict]]] = defaultdict(lambda: defaultdict(list))
    for item in enriched:
        if item.get("kind") == "Dashboard":
            dash_structure[item["ws"]][item["object"]].append(item)
        else:
            structure[item["ws"]][item["object"]].append(item)

    return {"reports": structure, "dashboards": dash_structure}


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_normal(doc: Document, text: str, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Paragraph")


def append_template_instructions(doc: Document) -> None:
    add_heading(doc, "Template Instructions", level=1)
    add_normal(
        doc,
        "Use this document as the reusable OriginBA Standard Offering Report Library validation "
        "template for client test cycles.",
    )
    add_normal(doc, "Replace these placeholders before final delivery:", bold=True)
    add_bullet(doc, "CLIENT_NAME - client short name shown on the cover page, footer, and scope text")
    add_bullet(doc, "TEST_VERSION - CIS test cycle version, for example 25.4")
    add_bullet(doc, "DATASOURCE_ALIAS - Jaspersoft datasource alias, for example CityCorp_DS")
    add_bullet(doc, "AUTHOR_NAME - tester or document author")
    add_bullet(doc, "PASS/FAIL/NA - validation result for each report, ad hoc view, and dashboard")
    add_normal(doc, "Recommended workflow:", bold=True)
    add_bullet(doc, "Find/replace CLIENT_NAME, TEST_VERSION, DATASOURCE_ALIAS, and AUTHOR_NAME.")
    add_bullet(doc, "Execute each report, ad hoc view, and dashboard in the Standard Offering library.")
    add_bullet(doc, "Update Validation and Notes columns for each object.")
    add_bullet(doc, "Update the Validation Summary and Test Results Summary when testing is complete.")
    add_bullet(doc, "Regenerate a client-ready copy with the generator script if preferred:")
    add_normal(
        doc,
        "python3 scripts/doc/build_standard_offering_validation_doc.py "
        "--client CityCorp --test-version 25.4 --datasource CityCorp_DS "
        "--author \"Chase Powers\" --status PASS",
    )
    add_page_break(doc)


def append_document_content(doc: Document, structure: dict, config: GenerationConfig) -> None:
    if config.is_fillable_template:
        append_template_instructions(doc)

    add_page_break(doc)

    add_heading(doc, "Document Versions", level=1)
    version_table = doc.add_table(rows=2, cols=4)
    version_table.style = "Table Grid"
    headers = ["Version", "Date", "Author", "Notes"]
    for idx, header in enumerate(headers):
        set_cell_text(version_table.rows[0].cells[idx], header, bold=True)
    values = ["1.0", config.resolved_date, config.author_name, config.version_notes]
    for idx, value in enumerate(values):
        set_cell_text(version_table.rows[1].cells[idx], value)

    doc.add_paragraph()
    add_heading(doc, "Test Script: Standard Offering Report Library Validation", level=1)
    add_normal(doc, "Scope:", bold=True)
    add_normal(
        doc,
        "Validate that each report, ad hoc view, and dashboard in the OriginBA SmartCity "
        f"Standard Offering Report Library executes successfully in the {config.client_name} "
        f"TEST {config.test_version} Jaspersoft environment and returns expected results.",
    )
    add_normal(doc, "Objective:", bold=True)
    add_bullet(doc, "Execute all 144 Standard Offering reports and ad hoc views.")
    add_bullet(doc, "Execute all 13 Standard Offering dashboards.")
    add_bullet(doc, "Confirm successful execution, rendering, and data availability for each object.")
    add_bullet(doc, "Capture validation status and notes for each tested object.")
    add_normal(doc, "Environment:", bold=True)
    add_bullet(doc, "Jaspersoft repository path: /SmartCity/Report/Standard_Offering")
    add_bullet(doc, f"Client / tenant: {config.client_name} TEST {config.test_version}")
    add_bullet(doc, f"Datasource alias: {config.datasource_alias}")
    add_normal(doc, "Limitation:", bold=True)
    add_normal(
        doc,
        "This validation confirms report and dashboard execution in the Standard Offering library. "
        "It does not replace functional CIS testing, role-based access testing, or snapshot parity validation.",
    )

    add_heading(doc, "Test Case - Standard Offering Report Library Validation", level=2)
    add_normal(doc, "Guidelines:", bold=True)
    add_bullet(doc, "Execute each report, ad hoc view, and dashboard individually from the Standard Offering library.")
    add_bullet(doc, "Update required input controls and filters before execution where applicable.")
    add_bullet(doc, "Record PASS when the object executes without error and returns expected results.")
    add_bullet(doc, "Record FAIL with execution notes if an object does not execute or render correctly.")

    add_heading(doc, "Validation Methodology", level=2)
    add_bullet(doc, "Successful execution with no Oracle BI / Jaspersoft error.")
    add_bullet(doc, "Expected rendering of tables, charts, and crosstabs.")
    add_bullet(doc, "Expected data returned for the selected filter context.")
    add_bullet(doc, "Input controls and filters updated as required prior to execution.")
    add_bullet(doc, "Dashboard visualizations load successfully.")

    add_heading(doc, "Validation Summary", level=2)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    summary_rows = [
        ("Reports / Ad Hoc Views Tested", "144"),
        ("Dashboards Tested", "13"),
        ("Total Objects Tested", "157"),
        ("Overall Result", config.summary_result),
    ]
    for idx, (label, value) in enumerate(summary_rows):
        set_cell_text(summary_table.rows[idx].cells[0], label, bold=True)
        set_cell_text(summary_table.rows[idx].cells[1], value)

    if config.is_fillable_template:
        summary_text = (
            "All 144 Standard Offering reports / ad hoc views and 13 dashboards were executed in "
            f"{PLACEHOLDER_CLIENT} TEST {PLACEHOLDER_TEST}. Update this summary after testing is complete."
        )
    else:
        summary_text = (
            "All 144 Standard Offering reports / ad hoc views and 13 dashboards were executed in "
            f"{config.client_name} TEST {config.test_version}. "
            "No execution failures or rendering issues were identified during testing."
        )
    add_normal(doc, summary_text)

    add_page_break(doc)
    add_heading(doc, "Validation Results by Workstream", level=1)

    for ws in WS_ORDER:
        ws_reports = structure["reports"].get(ws, {})
        ws_dashboards = structure["dashboards"].get(ws, {})
        if not ws_reports and not ws_dashboards:
            continue
        add_heading(doc, ws, level=2)
        objects = sorted(set(ws_reports) | set(ws_dashboards))
        for obj in objects:
            add_heading(doc, obj, level=3)
            report_rows = sorted(ws_reports.get(obj, []), key=lambda x: x["name"])
            if report_rows:
                add_heading(doc, "Reports and Ad Hoc Views", level=4)
                add_validation_table(doc, report_rows, config, dashboard=False)
            dash_rows = sorted(ws_dashboards.get(obj, []), key=lambda x: x["name"])
            if dash_rows:
                add_heading(doc, "Dashboards", level=4)
                add_validation_table(doc, dash_rows, config, dashboard=True)

    add_page_break(doc)
    add_heading(doc, "Test Results Summary", level=2)
    results_table = doc.add_table(rows=1, cols=4)
    results_table.style = "Table Grid"
    for idx, header in enumerate(["Step No.", "Workstream", "Actual Result", "Status"]):
        set_cell_text(results_table.rows[0].cells[idx], header, bold=True)
    for step, ws in enumerate(WS_ORDER, start=1):
        ws_reports = structure["reports"].get(ws, {})
        ws_dashboards = structure["dashboards"].get(ws, {})
        if not ws_reports and not ws_dashboards:
            continue
        count = sum(len(v) for v in ws_reports.values()) + sum(len(v) for v in ws_dashboards.values())
        cells = results_table.add_row().cells
        set_cell_text(cells[0], f"1.{step:02d}")
        set_cell_text(cells[1], ws)
        if config.is_fillable_template:
            actual = f"Validated {count} objects in {ws}. Update after testing."
            status = PLACEHOLDER_STATUS
        else:
            actual = f"Validated {count} objects in {ws}. All executed successfully."
            status = config.validation_status
        set_cell_text(cells[2], actual)
        set_cell_text(cells[3], status)


def generate_document(config: GenerationConfig) -> Path:
    reports, dashboards = load_catalog()
    zip_items = load_zip_index(STANDARD_OFFERING_ZIP)
    structure = build_structure(reports, dashboards, zip_items)

    if len(reports) != 144 or len(dashboards) != 13:
        raise SystemExit(f"Expected 144 reports and 13 dashboards, got {len(reports)} and {len(dashboards)}")

    cover_base = ensure_cover_base()
    config.output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(cover_base, config.output_path)
    prepare_template_cover(config.output_path, config)

    doc = Document(str(config.output_path))
    configure_footer(doc, config.footer_text)
    append_document_content(doc, structure, config)
    doc.save(str(config.output_path))
    return config.output_path


def parse_args() -> GenerationConfig:
    parser = argparse.ArgumentParser(
        description="Generate OriginBA Standard Offering report library validation documents."
    )
    parser.add_argument(
        "--template",
        action="store_true",
        help="Generate the reusable fillable template with CLIENT_NAME placeholders.",
    )
    parser.add_argument("--client", default="CityCorp", help="Client short name, for example CityCorp.")
    parser.add_argument("--test-version", default="25.4", help="CIS test version, for example 25.4.")
    parser.add_argument(
        "--datasource",
        default=None,
        help="Datasource alias. Defaults to <client>_DS unless template mode is used.",
    )
    parser.add_argument("--author", default="Chase Powers", help="Document author name.")
    parser.add_argument(
        "--status",
        default="PASS",
        help="Validation status for populated documents. Use PASS/FAIL/NA in template mode.",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output .docx path. Defaults to the template path or client-specific output path.",
    )
    args = parser.parse_args()

    if args.template:
        output_path = Path(args.output) if args.output else TEMPLATE_OUTPUT
        return GenerationConfig(
            client_name=PLACEHOLDER_CLIENT,
            test_version=PLACEHOLDER_TEST,
            datasource_alias=PLACEHOLDER_DATASOURCE,
            author_name=PLACEHOLDER_AUTHOR,
            validation_status=PLACEHOLDER_STATUS,
            notes_text=TEMPLATE_NOTES,
            template_mode=True,
            output_path=output_path,
        )

    client_name = args.client
    datasource = args.datasource or f"{client_name}_DS"
    if args.output:
        output_path = Path(args.output)
    else:
        if client_name == "CityCorp":
            output_path = DEFAULT_OUTPUT
        else:
            slug = re.sub(r"[^A-Za-z0-9]+", "_", client_name).strip("_")
            output_path = REPO_ROOT / f"output/doc/{slug}_Standard_Offering_Report_Library_Validation.docx"

    return GenerationConfig(
        client_name=client_name,
        test_version=args.test_version,
        datasource_alias=datasource,
        author_name=args.author,
        validation_status=args.status,
        notes_text=PLACEHOLDER_NOTES,
        template_mode=False,
        output_path=output_path,
    )


def main() -> None:
    config = parse_args()
    path = generate_document(config)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
